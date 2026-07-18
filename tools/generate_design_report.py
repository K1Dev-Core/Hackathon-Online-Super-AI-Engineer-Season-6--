#!/usr/bin/env python3
"""Build Thai design report by cloning retained Design Report reference DOCX."""

from __future__ import annotations

import hashlib
import re
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
SKILL_DIR = Path("/Users/k1god/.codex/plugins/cache/openai-curated-remote/openai-templates/0.1.0/skills/artifact-template-design-report")
REFERENCE = SKILL_DIR / "assets" / "reference.docx"
SOURCE = ROOT / "docs" / "SuperAI6_IoT_Attack_Detection_Design_Report_TH.md"
OUTPUT = ROOT / "docs" / "SuperAI6_IoT_Attack_Detection_Design_Report_TH.docx"
HERO = ROOT / "docs" / "assets" / "design_report_hero.png"

FONT = "Tahoma"
MONO = "Courier New"
BLACK = "000000"
WHITE = "FFFFFF"
CYAN = "2DD4BF"
CYAN_DARK = "0E8F87"
NAVY = "081522"
NAVY_2 = "10263A"
LIME = "B7F34A"
ORANGE = "FF8A4C"
LIGHT = "F2F6F7"
TEXT = "17212B"
MUTED = "596A77"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, value: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None, mono: bool = False) -> None:
    face = MONO if mono else FONT
    run.font.name = face
    run._element.rPr.rFonts.set(qn("w:eastAsia"), face)
    run._element.rPr.rFonts.set(qn("w:cs"), face)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text: str, *, size: float | None = None, color: str = TEXT, bold: bool = False) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    for part in token_re.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=(size - 0.5 if size else 9), bold=False, color=CYAN_DARK, mono=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=bold, color=color)


def make_hero(path: Path) -> None:
    width, height = 1800, 1740
    image = Image.new("RGB", (width, height), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    for x in range(0, width, 90):
        draw.line((x, 0, x, height), fill=(18, 42, 60), width=1)
    for y in range(0, height, 90):
        draw.line((0, y, width, y), fill=(18, 42, 60), width=1)
    draw.polygon([(0, 1180), (1800, 680), (1800, 1740), (0, 1740)], fill=(10, 39, 51))
    draw.polygon([(950, 0), (1800, 0), (1800, 900), (1300, 1180)], fill=(14, 62, 68))
    nodes = [
        (190, 330), (530, 190), (840, 410), (1160, 210), (1510, 410),
        (330, 780), (700, 720), (1080, 760), (1440, 890),
        (240, 1280), (650, 1190), (1020, 1390), (1510, 1290),
    ]
    edges = [(0,1),(1,2),(2,3),(3,4),(0,5),(2,6),(3,7),(4,8),(5,6),(6,7),(7,8),(5,9),(6,10),(7,11),(8,12),(9,10),(10,11),(11,12)]
    for a, b in edges:
        draw.line((*nodes[a], *nodes[b]), fill=(45, 212, 191), width=5)
    for idx, (x, y) in enumerate(nodes):
        radius = 18 if idx % 3 else 28
        fill = (183, 243, 74) if idx in (3, 8, 11) else (45, 212, 191)
        draw.ellipse((x-radius, y-radius, x+radius, y+radius), fill=fill, outline=(232, 255, 250), width=2)
    for offset in range(0, 420, 80):
        draw.rounded_rectangle((1050 + offset, 1510, 1100 + offset, 1545), radius=12, fill=(255, 138, 76))
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, quality=96)


def replace_package_image(docx_path: Path, member: str, image_path: Path) -> None:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        temp_path = Path(handle.name)
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        names = source.namelist()
        if member not in names:
            raise RuntimeError(f"Template image member missing: {member}")
        for item in source.infolist():
            data = image_path.read_bytes() if item.filename == member else source.read(item.filename)
            target.writestr(item, data)
    temp_path.replace(docx_path)


def clear_body_after_cover(doc: Document) -> None:
    body = doc._body._element
    children = list(body)
    first_break = None
    for index, child in enumerate(children):
        if child.tag == qn("w:p") and child.find("./w:pPr/w:sectPr", namespaces=child.nsmap) is not None:
            first_break = index
            break
    if first_break is None:
        raise RuntimeError("Reference first-section break not found")
    for child in children[first_break + 1:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document) -> None:
    for style_name in ("normal", "Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
    doc.styles["normal"].font.size = Pt(10)
    doc.styles["normal"].font.color.rgb = RGBColor.from_string(TEXT)
    doc.styles["normal"].paragraph_format.line_spacing = 1.12
    doc.styles["normal"].paragraph_format.space_after = Pt(5)


def configure_cover(doc: Document) -> None:
    title = next(paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "Report title")
    title.clear()
    title.style = doc.styles["Title"]
    run = title.add_run("The IoT Attack Detection\nChallenge")
    set_run_font(run, size=31, bold=True, color=BLACK)

    table = doc.tables[0]
    left = table.cell(0, 0)
    right = table.cell(0, 2)
    left.text = "การออกแบบ detector สำหรับ Normal-only data\nด้วย Protocol Rules, Stream Structure และ PCAP Evidence"
    right.text = "จัดทำโดย 610686-วชิรวิทย์\nกรกฎาคม 2569"
    for cell in (left, right):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=9.5, color=MUTED)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value or "-") for value in values):
            rows.append(values)
        index += 1
    return rows, index


def add_template_table(doc: Document, rows: list[list[str]], template_table) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    if template_table._tbl.tblPr is not None:
        table._tbl.replace(table._tbl.tblPr, deepcopy(template_table._tbl.tblPr))
    for row_index, values in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_fill(cell, BLACK if row_index == 0 else (LIGHT if row_index % 2 == 0 else WHITE))
            set_cell_margins(cell, 95)
            paragraph = cell.paragraphs[0]
            paragraph.style = doc.styles["normal"]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(
                paragraph,
                values[column_index] if column_index < len(values) else "",
                size=8.1,
                color=WHITE if row_index == 0 else TEXT,
                bold=row_index == 0,
            )
    set_repeat_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_fill(cell, NAVY_2)
    set_cell_margins(cell, 140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(code.rstrip())
    set_run_font(run, size=8, color=WHITE, mono=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc: Document, text: str, callout_ppr) -> None:
    paragraph = doc.add_paragraph()
    if callout_ppr is not None:
        paragraph._p.insert(0, deepcopy(callout_ppr))
    paragraph.clear()
    add_inline(paragraph, text, size=10, color=TEXT, bold=True)


def add_contents(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    add_inline(paragraph, "สารบัญ", size=17, color=BLACK, bold=True)
    for line in lines:
        if line.startswith("# ") and not line.startswith("# The IoT"):
            paragraph = doc.add_paragraph(style="normal")
            paragraph.paragraph_format.left_indent = Pt(8)
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline(paragraph, line[2:].strip(), size=10, color=MUTED)
    doc.add_page_break()


def add_markdown_body(doc: Document, lines: list[str], template_table, callout_ppr) -> None:
    start = next(i for i, line in enumerate(lines) if line.startswith("# บทสรุปผู้บริหาร"))
    index = start
    paragraph_buffer: list[str] = []
    h1_count = 0

    def flush() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        paragraph = doc.add_paragraph(style="normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(paragraph, " ".join(part.strip() for part in paragraph_buffer), size=10, color=TEXT)
        paragraph_buffer = []

    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            flush()
            index += 1
            continue
        if stripped.startswith("```"):
            flush()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            add_code_block(doc, "\n".join(code_lines))
            continue
        if stripped.startswith("|"):
            flush()
            rows, index = parse_table(lines, index)
            add_template_table(doc, rows, template_table)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush()
            level = len(heading.group(1))
            title = heading.group(2)
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            if level == 1:
                paragraph.paragraph_format.page_break_before = h1_count > 0
                h1_count += 1
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(title)
            size = 24 if level == 1 else 15 if level == 2 else 12
            set_run_font(run, size=size, bold=level != 1, color=BLACK)
            index += 1
            continue
        if stripped.startswith("> "):
            flush()
            add_callout(doc, stripped[2:], callout_ppr)
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        ordered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if bullet or ordered:
            flush()
            paragraph = doc.add_paragraph(style="normal")
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.first_line_indent = Pt(-12)
            prefix = "• " if bullet else f"{ordered.group(1)}. "
            add_inline(paragraph, prefix + (bullet.group(1) if bullet else ordered.group(2)), size=9.8, color=TEXT)
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush()


def build() -> None:
    reference_hash_before = sha256(REFERENCE)
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    make_hero(HERO)
    doc = Document(REFERENCE)
    template_table = doc.tables[1]
    callout = next((paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("Key takeaway.")), None)
    callout_ppr = deepcopy(callout._p.pPr) if callout is not None and callout._p.pPr is not None else None

    configure_styles(doc)
    configure_cover(doc)
    clear_body_after_cover(doc)
    add_contents(doc, source_lines)
    add_markdown_body(doc, source_lines, template_table, callout_ppr)

    properties = doc.core_properties
    properties.title = "The IoT Attack Detection Challenge — Design Report"
    properties.subject = "Super AI Engineer Season 6"
    properties.author = "610686-วชิรวิทย์"
    properties.keywords = "IoT, MQTT, anomaly detection, design report, PCAP"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    replace_package_image(OUTPUT, "word/media/image1.png", HERO)
    reference_hash_after = sha256(REFERENCE)
    if reference_hash_before != reference_hash_after:
        raise RuntimeError("Retained reference changed")
    print(f"reference_sha256={reference_hash_after}")
    print(f"output={OUTPUT}")


if __name__ == "__main__":
    build()

