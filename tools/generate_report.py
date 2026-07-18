#!/usr/bin/env python3
"""Generate styled Thai DOCX report from canonical Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SuperAI6_IoT_Attack_Detection_Report_TH.md"
OUTPUT = ROOT / "docs" / "SuperAI6_IoT_Attack_Detection_Report_TH.docx"
ASSET_DIR = ROOT / "docs" / "assets"
COVER = ASSET_DIR / "report_cover.png"

NAVY = "0B172A"
NAVY_2 = "13243A"
CYAN = "2DD4BF"
LIME = "B7F34A"
ORANGE = "FF8A4C"
LIGHT = "F4F7F8"
MID = "D8E2E7"
TEXT = "17212B"
MUTED = "5C6B78"
THAI_FONT = "Tahoma"
MONO_FONT = "Courier New"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("หน้า ")
    run.font.name = THAI_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_inline(paragraph, text: str, *, color: str = TEXT, size: float | None = None) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    for part in token_re.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt((size or 10.5) - 0.5)
            run.font.color.rgb = rgb("007A75")
        else:
            run = paragraph.add_run(part)
        run.font.name = run.font.name or THAI_FONT
        run.font.color.rgb = rgb(color)
        if size is not None and not (part.startswith("`") and part.endswith("`")):
            run.font.size = Pt(size)


def make_cover(path: Path) -> None:
    width, height = 1800, 900
    image = Image.new("RGB", (width, height), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    for x in range(0, width, 90):
        draw.line((x, 0, x, height), fill=(20, 43, 67), width=1)
    for y in range(0, height, 90):
        draw.line((0, y, width, y), fill=(20, 43, 67), width=1)
    draw.polygon([(1130, 0), (1800, 0), (1800, 900), (1450, 900)], fill=(16, 49, 69))
    draw.polygon([(1450, 0), (1800, 0), (1800, 570)], fill=(31, 87, 93))
    nodes = [(1120, 180), (1390, 120), (1580, 300), (1280, 430), (1660, 540), (1420, 690), (1110, 720)]
    edges = [(0, 1), (1, 2), (0, 3), (3, 2), (3, 4), (3, 5), (5, 4), (5, 6)]
    for a, b in edges:
        draw.line((*nodes[a], *nodes[b]), fill=(45, 212, 191), width=5)
    for idx, (x, y) in enumerate(nodes):
        radius = 15 if idx % 2 else 22
        fill = (183, 243, 74) if idx in (2, 5) else (45, 212, 191)
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=fill)
    draw.rounded_rectangle((105, 625, 860, 730), radius=22, fill=(45, 212, 191))
    draw.rounded_rectangle((105, 755, 650, 815), radius=18, outline=(183, 243, 74), width=3)
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, quality=95)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = THAI_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in (
        ("Title", 28, NAVY),
        ("Subtitle", 15, MUTED),
        ("Heading 1", 20, NAVY),
        ("Heading 2", 15, "087F79"),
        ("Heading 3", 12, NAVY_2),
    ):
        style = styles[name]
        style.font.name = THAI_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("SUPER AI ENGINEER SEASON 6  /  IoT ATTACK DETECTION")
    r.font.name = THAI_FONT
    r.font.size = Pt(7.5)
    r.font.bold = True
    r.font.color.rgb = rgb("0B8F88")
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    doc.add_picture(str(COVER), width=Inches(6.85))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    add_inline(p, "THE IoT ATTACK DETECTION CHALLENGE", color=NAVY, size=27)
    p.runs[0].bold = True
    p = doc.add_paragraph()
    add_inline(p, "การตรวจจับการโจมตีบนเครือข่าย IoT ด้วย Protocol Rules, Stream Structure และ PCAP Evidence", color="087F79", size=15)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_inline(p, "610686-วชิรวิทย์  |  Super AI Engineer Season 6", color=MUTED, size=11)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(8.2)
    table.columns[1].width = Cm(8.2)
    values = [("PUBLIC F1", "0.96267", CYAN), ("PRIVATE F1", "0.97198", LIME)]
    for idx, (label, value, accent) in enumerate(values):
        cell = table.cell(0, idx)
        set_cell_fill(cell, NAVY_2)
        set_cell_margins(cell, 180, 220, 180, 220)
        p = cell.paragraphs[0]
        add_inline(p, label + "\n", color="AFC3CF", size=9)
        run = p.add_run(value)
        run.font.name = THAI_FONT
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = rgb(accent)
    doc.add_page_break()


def add_toc(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("สารบัญ").font.name = THAI_FONT
    for line in lines:
        if line.startswith("# ") and not line.startswith("# รายงาน"):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.2)
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, text, color=TEXT, size=10.5)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        values = [x.strip() for x in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value or "-") for value in values):
            rows.append(values)
        idx += 1
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, values in enumerate(rows):
        for col_idx in range(columns):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_fill(cell, NAVY_2 if row_idx == 0 else ("EEF4F5" if row_idx % 2 else "FFFFFF"))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(
                p,
                values[col_idx] if col_idx < len(values) else "",
                color="FFFFFF" if row_idx == 0 else TEXT,
                size=8.5,
            )
            for run in p.runs:
                run.bold = row_idx == 0
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_fill(cell, NAVY_2)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code.rstrip())
    run.font.name = MONO_FONT
    run.font.size = Pt(8.2)
    run.font.color.rgb = rgb("DCE9EE")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_quote(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_fill(cell, "E5F8F5")
    set_cell_margins(cell, 140, 220, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, color="075E59", size=10.5)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    make_cover(COVER)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc, lines)

    start = next(i for i, line in enumerate(lines) if line.startswith("# 1."))
    idx = start
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, " ".join(x.strip() for x in paragraph_buffer), size=10.5)
            paragraph_buffer = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            idx += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            code_lines: list[str] = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            idx += 1
            add_code_block(doc, "\n".join(code_lines))
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            title = heading.group(2)
            if level == 1 and doc.paragraphs:
                doc.add_page_break()
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, title, color=NAVY if level == 1 else ("087F79" if level == 2 else NAVY_2))
            idx += 1
            continue
        if stripped.startswith("> "):
            flush_paragraph()
            add_quote(doc, stripped[2:])
            idx += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or ordered:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            add_inline(p, (bullet or ordered).group(1), size=10.2)
            idx += 1
            continue
        paragraph_buffer.append(stripped)
        idx += 1

    flush_paragraph()
    props = doc.core_properties
    props.title = "The IoT Attack Detection Challenge — รายงานโครงการ"
    props.subject = "Super AI Engineer Season 6"
    props.author = "610686-วชิรวิทย์"
    props.keywords = "IoT, MQTT, anomaly detection, PCAP, F1-score"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"output={OUTPUT}")


if __name__ == "__main__":
    build()
